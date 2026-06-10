"""Tests for LAP old/new format migration onto the current template."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from timetable.bundle_paths import resource_path
from timetable.io.lap_docx import build_export_lap, update_lap_lecturer
from timetable.io.lap_docx_utils import set_cell_text
from timetable.io.lap_migrate import extract_lap, migrate_lap_to_template

_OLD_LAP = Path(
    "/Users/tomdrummond/Documents/apps/FormFiller-Innerrer/uploads/"
    "20260528-114128_ICTNWK424_ICTTEN434_ICTNWK429_LAP_EP_FT_S12023.docx"
)
_TEMPLATE = resource_path("templates", "lap_template.docx")


@pytest.fixture
def old_lap_bytes() -> bytes:
    if not _OLD_LAP.is_file():
        pytest.skip("sample old LAP not available")
    return _OLD_LAP.read_bytes()


def test_extract_old_lap_has_units_and_sessions(old_lap_bytes):
    data = extract_lap(old_lap_bytes)
    assert len(data["units"]) == 3
    assert len(data["assessments"]) == 4
    assert len(data["sessions"]) == 21
    assert data["delivery_location"]
    assert data["sessions"][0]["hrs"] == "3"
    assert data["sessions"][1]["hrs"] == "3"


def test_migrate_zero_padded_session_numbers(old_lap_bytes):
    """Session labels like 01–09 must match template rows 1–9."""
    migrated = migrate_lap_to_template(old_lap_bytes)
    roundtrip = extract_lap(migrated)
    assert len(roundtrip["sessions"]) >= 20


def test_migrate_cyber_security_awareness_lap_all_sessions():
    path = Path("/Users/tomdrummond/Desktop/Cyber Security Awareness LAP.docx")
    if not path.is_file():
        pytest.skip("Cyber Security Awareness LAP sample not available")
    data = extract_lap(path.read_bytes())
    assert len(data["sessions"]) == 20
    assert data["sessions"][0]["session"] in ("01", "1")
    migrated = migrate_lap_to_template(path.read_bytes())
    roundtrip = extract_lap(migrated)
    topics = [s["topic"] for s in roundtrip["sessions"] if (s.get("topic") or "").strip()]
    assert len(topics) == 20

    from timetable.io.lap_docx_utils import distinct_cells

    doc = Document(BytesIO(migrated))
    table = doc.tables[5]
    training_row = next(
        r for r in table.rows if "total training hours" in (r.cells[0].text or "").lower()
    )
    unit_row = next(
        r
        for r in table.rows
        if "total amount of training for this unit" in (r.cells[0].text or "").lower()
    )
    training_cells = distinct_cells(training_row)
    unit_cells = distinct_cells(unit_row)
    assert training_cells[1].text.strip() == "60"
    assert training_cells[3].text.strip() == "18"
    assert unit_cells[1].text.strip() == "78hrs"


def test_migrate_old_lap_uses_assessment_task_header(old_lap_bytes):
    migrated = migrate_lap_to_template(old_lap_bytes)
    doc = Document(BytesIO(migrated))
    table = next(
        t
        for t in doc.tables
        if any("assessment task" in (r.cells[0].text or "").lower() for r in t.rows)
    )
    task_rows = [r for r in table.rows if (r.cells[0].text or "").lower().startswith("assessment task")]
    assert len(task_rows) >= 4
    assert task_rows[0].cells[1].text.strip()


def test_migrate_old_lap_session_hours_are_black(old_lap_bytes):
    migrated = migrate_lap_to_template(old_lap_bytes)
    doc = Document(BytesIO(migrated))
    for t in doc.tables:
        for r in t.rows:
            if r.cells[0].text.strip() in ("1", "2"):
                run = r.cells[1].paragraphs[0].runs[0]
                assert run.font.color.rgb == RGBColor(0, 0, 0)
                assert r.cells[1].text.strip() == "3"
                if r.cells[0].text.strip() == "2":
                    return


def test_set_cell_text_forces_aptos_12():
    if not _TEMPLATE.is_file():
        pytest.skip("lap_template.docx not available")
    doc = Document(_TEMPLATE)
    cell = doc.tables[0].rows[0].cells[1]
    set_cell_text(cell, "Example value")
    run = cell.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor(0, 0, 0)
    assert run.font.name == "Aptos"
    assert run.font.size.pt == 12


def test_set_cell_text_forces_black():
    if not _TEMPLATE.is_file():
        pytest.skip("lap_template.docx not available")
    doc = Document(_TEMPLATE)
    cell = doc.tables[0].rows[0].cells[1]
    set_cell_text(cell, "Example value")
    run = cell.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor(0, 0, 0)


def test_migrated_lap_written_fields_use_aptos_12(old_lap_bytes):
    migrated = migrate_lap_to_template(old_lap_bytes)
    doc = Document(BytesIO(migrated))
    checked = 0
    for t in doc.tables:
        for r in t.rows:
            for cell in r.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        text = (run.text or "").strip()
                        if not text or text.lower().startswith("total"):
                            continue
                        if run.font.name != "Aptos":
                            continue
                        if run.font.size is None:
                            continue
                        assert run.font.size.pt == 12
                        checked += 1
                        if checked >= 20:
                            return
    assert checked >= 20


def test_build_export_lap_applies_delivery_period():
    if not _TEMPLATE.is_file():
        pytest.skip("lap_template.docx not available")
    exported = build_export_lap(
        _TEMPLATE.read_bytes(),
        {"name": "Jane Smith", "email": "jane.smith@nmtafe.wa.edu.au"},
        delivery_period="2026 Semester 1 (20 weeks)",
    )
    from io import BytesIO

    from timetable.io.lap_migrate import extract_lap

    data = extract_lap(BytesIO(exported))
    assert data["delivery_period"] == "2026 Semester 1 (20 weeks)"


def test_build_export_lap_applies_multiple_lecturers():
    if not _TEMPLATE.is_file():
        pytest.skip("lap_template.docx not available")
    exported = build_export_lap(
        _TEMPLATE.read_bytes(),
        [
            {
                "name": "Jane Smith",
                "phone": "N/A",
                "email": "jane.smith@nmtafe.wa.edu.au",
                "contact_times": "regular business hours",
            },
            {
                "name": "Bob Jones",
                "phone": "N/A",
                "email": "bob.jones@nmtafe.wa.edu.au",
                "contact_times": "regular business hours",
            },
        ],
    )
    from timetable.io.lap_docx import _extract_lecturers

    doc = Document(BytesIO(exported))
    lecturers = _extract_lecturers(doc)
    assert len(lecturers) == 2
    assert lecturers[0]["name"] == "Jane Smith"
    assert lecturers[1]["name"] == "Bob Jones"


def test_update_lap_lecturer_still_accepts_single_dict():
    if not _TEMPLATE.is_file():
        pytest.skip("lap_template.docx not available")
    exported = update_lap_lecturer(
        _TEMPLATE.read_bytes(),
        {"name": "Jane Smith", "email": "jane.smith@nmtafe.wa.edu.au"},
    )
    from timetable.io.lap_docx import _extract_lecturer

    doc = Document(BytesIO(exported))
    assert _extract_lecturer(doc)["name"] == "Jane Smith"
