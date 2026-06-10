"""Learning & Assessment Plan (.docx) read/update for timetable lecturer sync."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from .lap_docx_utils import (
    cell_text,
    copy_row_after,
    distinct_cells,
    find_table,
    footer_codes,
    norm,
    set_cell_text,
)

LAP_FOOTER_CODE = "F122A14"
_LECTURER_KEYS = ("name", "phone", "email", "contact_times", "campus")


def is_lap_document(path: str | Path | BytesIO) -> bool:
    """True when the document footer matches the LAP form code."""
    doc = Document(path)
    return LAP_FOOTER_CODE in footer_codes(doc)


def _lecturer_table(doc: Document):
    return find_table(
        doc,
        lambda t: any(
            "student to supply" in norm(cell_text(distinct_cells(r)[0])) for r in t.rows
        ),
    )


def _extract_lecturers(doc: Document) -> list[dict[str, str]]:
    table = _lecturer_table(doc)
    if table is None:
        return []
    rows = table.rows
    header_idx = next(
        (
            i
            for i, row in enumerate(rows)
            if "lecturer name" in norm(cell_text(distinct_cells(row)[0]))
        ),
        None,
    )
    if header_idx is None:
        return []
    lecturers: list[dict[str, str]] = []
    for row in rows[header_idx + 1 :]:
        value_cells = row.cells
        getv = lambda idx: cell_text(value_cells[idx]) if idx < len(value_cells) else ""
        entry = {key: getv(i) for i, key in enumerate(_LECTURER_KEYS)}
        if any(entry.values()):
            lecturers.append(entry)
    return lecturers


def _extract_lecturer(doc: Document) -> dict[str, str]:
    lecturers = _extract_lecturers(doc)
    if lecturers:
        return lecturers[0]
    return {key: "" for key in _LECTURER_KEYS}


def _fill_lecturer_rows(doc: Document, lecturers: list[dict[str, str]]) -> bool:
    if not lecturers:
        return False
    table = _lecturer_table(doc)
    if table is None:
        return False
    rows = list(table.rows)
    header_idx = next(
        i
        for i, row in enumerate(rows)
        if "lecturer name" in norm(cell_text(distinct_cells(row)[0]))
    )

    def fill_row(row, lec: dict[str, str]) -> None:
        value_cells = row.cells
        for idx, key in enumerate(_LECTURER_KEYS):
            value = (lec.get(key) or "").strip()
            if value and idx < len(value_cells):
                set_cell_text(value_cells[idx], value)

    first_data_row = rows[header_idx + 1]
    fill_row(first_data_row, lecturers[0])
    prev = first_data_row
    for lec in lecturers[1:]:
        prev = copy_row_after(table, prev)
        fill_row(prev, lec)
    return True


def _session_table(doc: Document):
    return find_table(
        doc,
        lambda t: any(
            norm(cell_text(r.cells[0])) == "session" and norm(cell_text(r.cells[1])) == "hrs"
            for r in t.rows
            if len(r.cells) >= 2
        ),
    )


def overlay_session_hours(doc: Document, hours_by_session: dict[int | str, str]) -> None:
    """Write contact hours into numbered session rows (e.g. 1 and 2 for double sessions)."""
    if not hours_by_session:
        return
    normalized = {str(k): v for k, v in hours_by_session.items()}
    table = _session_table(doc)
    if table is None:
        return
    rows = list(table.rows)
    subheader = next(
        i
        for i, r in enumerate(rows)
        if len(r.cells) >= 2
        and norm(cell_text(r.cells[0])) == "session"
        and norm(cell_text(r.cells[1])) == "hrs"
    )
    for row in rows[subheader + 1 :]:
        num = norm(cell_text(row.cells[0]))
        if num.startswith("total"):
            break
        if num in normalized:
            set_cell_text(row.cells[1], normalized[num])


def _merge_lecturers(
    uploaded: list[dict[str, str]],
    timetable: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Apply timetable contact fields; keep campus from upload when timetable has none."""
    campuses = [u.get("campus", "").strip() for u in uploaded if u.get("campus", "").strip()]
    default_campus = campuses[0] if campuses else ""
    merged: list[dict[str, str]] = []
    for i, lec in enumerate(timetable):
        campus = (lec.get("campus") or "").strip()
        if not campus:
            campus = uploaded[i].get("campus", "").strip() if i < len(uploaded) else default_campus
        merged.append({**lec, "campus": campus})
    return merged


def build_export_lap(
    source: bytes | str | Path,
    lecturers: list[dict[str, str]] | dict[str, str],
    *,
    session_hours: dict[int | str, str] | None = None,
    delivery_period: str | None = None,
) -> bytes:
    """Migrate uploaded LAP content onto the current template, then apply timetable lecturers."""
    from .lap_migrate import migrate_lap_to_template, set_delivery_period, update_session_hour_totals

    if isinstance(lecturers, dict):
        lecturers = [lecturers]

    migrated = migrate_lap_to_template(source)
    doc = Document(BytesIO(migrated))
    uploaded = _extract_lecturers(doc)
    merged = _merge_lecturers(uploaded, lecturers)
    if not _fill_lecturer_rows(doc, merged):
        raise ValueError("Could not find lecturer details table in LAP document")
    if session_hours:
        overlay_session_hours(doc, session_hours)
    if delivery_period:
        set_delivery_period(doc, delivery_period)
    update_session_hour_totals(doc)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def update_lap_lecturer(
    source: bytes | str | Path,
    lecturer: dict[str, str] | list[dict[str, str]],
) -> bytes:
    """Patch lecturer field(s) in an existing LAP."""
    if isinstance(lecturer, dict):
        lecturers = [lecturer]
    else:
        lecturers = lecturer
    bio_in = BytesIO(source) if isinstance(source, (bytes, bytearray)) else None
    doc = Document(bio_in if bio_in is not None else source)
    uploaded = _extract_lecturers(doc)
    merged = _merge_lecturers(uploaded, lecturers)
    if not _fill_lecturer_rows(doc, merged):
        raise ValueError("Could not find lecturer details table in LAP document")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
