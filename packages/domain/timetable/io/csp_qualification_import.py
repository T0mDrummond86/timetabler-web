"""Import qualifications + classes from NMTAFE Curriculum Structure Package (CSP) .docx files.

Each table in the document is one qualification stage (e.g. Semester 1 / Semester 2).
From each table we read only:
  - column 0: class name and weekly contact hours (``Name | 2hrs``)
  - column 2: unit code (TPN — the column after SIN)

Multi-unit classes may use a vertically merged first column; continuation rows inherit
the current class name and hours.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from ..core.models import Course, Qualification, Unit, UnitQualification
from ..core.qualification_schedule import SCHEDULE_PERIOD_DAY, replace_qualification_time_windows
from ..core.unit_brackets import apply_unit_bracket_fields_from_names, normalize_component_codes_commas
from .qualification_import import (
    QualImportReport,
    _create_kwargs,
    _parse_running_time,
    _scoped_filter_by,
    _text,
)

_CLASS_CELL_RE = re.compile(
    r"^(.+?)\s*(?:\||\n)\s*(\d+(?:\.\d+)?)\s*hrs?\s*$",
    re.IGNORECASE | re.DOTALL,
)
_STAGE_PREFIX_RE = re.compile(
    r"^(semester\s+\d+|stage\s+\d+|year\s+\d+|term\s+\d+)",
    re.IGNORECASE,
)


@dataclass
class CspClass:
    name: str
    hours: float | None = None
    unit_codes: list[str] = field(default_factory=list)


@dataclass
class CspStage:
    qualification_name: str
    stage_label: str | None
    classes: list[CspClass]


def _iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _cell_text(cell) -> str:
    return (cell.text or "").strip()


def _is_vmerge_continue(cell) -> bool:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return False
    vm = tc_pr.find(qn("w:vMerge"))
    if vm is None:
        return False
    return vm.get(qn("w:val")) != "restart"


def _logical_class_cell_text(row) -> str:
    if not row.cells:
        return ""
    cell = row.cells[0]
    if _is_vmerge_continue(cell):
        return ""
    return _cell_text(cell)


def _stage_label_from_paragraph(text: str) -> str | None:
    s = _text(text)
    if s is None:
        return None
    for sep in (" – ", " - ", "—", "–"):
        if sep in s:
            s = s.split(sep, 1)[0].strip()
            break
    m = _STAGE_PREFIX_RE.match(s)
    if not m:
        return None
    words = m.group(1).split()
    return " ".join(w.capitalize() if w.isalpha() else w for w in words)


def _parse_class_cell(text: str) -> tuple[str, float | None]:
    s = (text or "").strip()
    if not s:
        return "", None
    m = _CLASS_CELL_RE.match(s)
    if m:
        return m.group(1).strip(), float(m.group(2))
    return s, None


def _is_csp_table(table: Table) -> bool:
    if not table.rows:
        return False
    header = table.rows[0]
    if len(header.cells) < 3:
        return False
    labels = [_cell_text(c).lower() for c in header.cells[:4]]
    return "sin" in labels and any("tpn" in lbl or lbl == "tpn" for lbl in labels)


def _parse_csp_table(table: Table) -> list[CspClass]:
    classes: list[CspClass] = []
    current: CspClass | None = None

    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue
        unit_code = _cell_text(row.cells[2]) if len(row.cells) > 2 else ""
        if not unit_code or unit_code.lower() in ("tpn", "uoc(s) being assessed in skill set"):
            continue

        raw_class = _logical_class_cell_text(row)
        if raw_class:
            name, hours = _parse_class_cell(raw_class)
            if not name:
                continue
            key = (name, hours)
            if current is None or (current.name, current.hours) != key:
                current = CspClass(name=name, hours=hours, unit_codes=[])
                classes.append(current)
        elif current is None:
            continue

        current.unit_codes.append(unit_code.strip())

    return classes


def _qualification_title_from_doc(doc: Document, path: Path | None) -> str:
    for para in doc.paragraphs:
        text = _text(para.text)
        if text and not text.lower().startswith("pathway code"):
            return text
    if path is not None:
        stem = path.stem
        if stem.upper().startswith("CSP_"):
            stem = stem[4:]
        return stem.replace("_", " ").strip()
    return "Imported qualification"


def _stage_qualification_name(base_title: str, stage_label: str | None, stage_index: int) -> str:
    if stage_label:
        return f"{base_title} – {stage_label}"
    if stage_index > 1:
        return f"{base_title} – Stage {stage_index}"
    return base_title


def extract_csp_qualification_stages(
    source: bytes | Path | str | Document,
    *,
    path_hint: Path | None = None,
) -> list[CspStage]:
    """Parse CSP tables into per-stage qualification payloads."""
    if hasattr(source, "tables") and hasattr(source, "element"):
        doc = source
        path = path_hint
    elif isinstance(source, (bytes, bytearray)):
        doc = Document(BytesIO(source))
        path = path_hint
    else:
        path = Path(source)
        doc = Document(path)

    base_title = _qualification_title_from_doc(doc, path)
    stages: list[CspStage] = []
    pending_stage_label: str | None = None
    stage_index = 0

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            label = _stage_label_from_paragraph(item.text)
            if label:
                pending_stage_label = label
            continue

        if not _is_csp_table(item):
            continue

        classes = _parse_csp_table(item)
        if not classes:
            pending_stage_label = None
            continue

        stage_index += 1
        qual_name = _stage_qualification_name(base_title, pending_stage_label, stage_index)
        stages.append(
            CspStage(
                qualification_name=qual_name,
                stage_label=pending_stage_label,
                classes=classes,
            )
        )
        pending_stage_label = None

    return stages


def import_qualifications_from_csp(
    session: Session,
    path: str | Path,
    *,
    timetable_session_id: int | None = None,
) -> QualImportReport:
    """Create/update qualifications and classes from a CSP .docx file."""
    rep = QualImportReport()
    path = Path(path)
    stages = extract_csp_qualification_stages(path)

    if not stages:
        rep.warnings.append(f"No CSP qualification tables found in {path.name}")
        return rep

    for stage in stages:
        qual_name = stage.qualification_name
        qual = _scoped_filter_by(
            session, Qualification, timetable_session_id, name=qual_name
        ).first()
        if qual is None:
            qual = Qualification(
                **_create_kwargs(
                    Qualification,
                    timetable_session_id,
                    name=qual_name,
                    num_groups=1,
                    schedule_period=SCHEDULE_PERIOD_DAY,
                )
            )
            session.add(qual)
            session.flush()
            replace_qualification_time_windows(session, qual)
            rep.qualifications_created += 1
            default_course_code = f"{qual.name} GrpA"
            existing_course = _scoped_filter_by(
                session, Course, timetable_session_id, code=default_course_code
            ).first()
            if existing_course is None:
                session.add(
                    Course(
                        **_create_kwargs(
                            Course,
                            timetable_session_id,
                            code=default_course_code,
                            qualification_id=qual.id,
                        )
                    )
                )
                rep.courses_created += 1
            elif existing_course.qualification_id != qual.id:
                existing_course.qualification_id = qual.id
        else:
            rep.qualifications_linked += 1

        for cls in stage.classes:
            storage_name = cls.name.strip()
            if not storage_name:
                continue

            component_codes = normalize_component_codes_commas(", ".join(cls.unit_codes))
            existing_unit = _scoped_filter_by(
                session, Unit, timetable_session_id, name=storage_name
            ).first()
            if existing_unit is None:
                unit = Unit(
                    **_create_kwargs(
                        Unit,
                        timetable_session_id,
                        name=storage_name,
                        component_codes=component_codes,
                    )
                )
                session.add(unit)
                session.flush()
                rep.classes_created += 1
            else:
                unit = existing_unit
                rep.classes_updated += 1

            if component_codes and not (unit.component_codes or "").strip():
                unit.component_codes = component_codes
            elif component_codes:
                merged = normalize_component_codes_commas(
                    f"{unit.component_codes}, {component_codes}"
                )
                if merged and merged != (unit.component_codes or "").strip():
                    unit.component_codes = merged

            if cls.hours is not None and not unit.length_slots:
                slots = _parse_running_time(str(cls.hours))
                if slots:
                    unit.length_slots = slots

            link = (
                session.query(UnitQualification)
                .filter_by(unit_id=unit.id, qualification_id=qual.id)
                .first()
            )
            if link is None:
                session.add(UnitQualification(unit_id=unit.id, qualification_id=qual.id))
                rep.class_qual_links_added += 1

    apply_unit_bracket_fields_from_names(session, timetable_session_id=timetable_session_id)
    session.commit()
    return rep
