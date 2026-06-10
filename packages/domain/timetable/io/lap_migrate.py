"""Learning & Assessment Plan (F122A14) migration — old or new format → current template."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from ..bundle_paths import resource_path
from .lap_docx_utils import (
    cell_text,
    copy_row_after,
    distinct_cells,
    find_row,
    find_table,
    norm,
    set_cell_text,
)

LAP_TEMPLATE_PATH = resource_path("templates", "lap_template.docx")


def _session_lookup_key(label: str) -> str:
    """Match session rows regardless of zero-padding (``01`` vs ``1``)."""
    text = (label or "").strip()
    if text.isdigit():
        return str(int(text))
    return norm(text)


def _info_value(doc: Document, *needles: str) -> str:
    for table in doc.tables:
        row = find_row(table, *needles)
        if row is not None:
            cells = distinct_cells(row)
            if len(cells) >= 2:
                return cell_text(cells[-1])
    return ""


def _extract_units(doc: Document) -> tuple[list[dict[str, str]], str]:
    table = find_table(
        doc,
        lambda t: any(
            norm(cell_text(distinct_cells(r)[0])) in ("national id", "national code")
            for r in t.rows
        ),
    )
    units: list[dict[str, str]] = []
    location = ""
    if table is None:
        return units, location

    started = False
    for row in table.rows:
        cells = distinct_cells(row)
        first = norm(cell_text(cells[0]))
        if first in ("national id", "national code"):
            started = True
            continue
        if not started:
            continue
        if "delivery location" in first:
            location = cell_text(cells[-1]) if len(cells) >= 2 else ""
            break
        if "you can access the full unit" in first:
            continue
        code = cell_text(cells[0])
        name = cell_text(cells[1]) if len(cells) >= 2 else ""
        if code:
            units.append({"code": code, "name": name})
    return units, location


def _lecturer_row_dict(value_cells) -> dict[str, str]:
    getv = lambda idx: cell_text(value_cells[idx]) if idx < len(value_cells) else ""
    return {
        "name": getv(0),
        "phone": getv(1),
        "email": getv(2),
        "contact_times": getv(3),
        "campus": getv(4),
    }


def _extract_lecturer_and_supply(doc: Document) -> dict:
    result = {
        "student_supply": "",
        "college_supply": "",
        "lecturer": {
            "name": "",
            "phone": "",
            "email": "",
            "contact_times": "",
            "campus": "",
        },
        "lecturers": [],
    }
    table = find_table(
        doc,
        lambda t: any(
            "student to supply" in norm(cell_text(distinct_cells(r)[0])) for r in t.rows
        ),
    )
    if table is None:
        return result

    rows = table.rows
    header_idx: int | None = None
    for i, row in enumerate(rows):
        cells = distinct_cells(row)
        first = norm(cell_text(cells[0]))
        if "student to supply" in first:
            result["student_supply"] = cell_text(cells[0])
        elif "college to supply" in first:
            result["college_supply"] = cell_text(cells[0])
        elif "lecturer name" in first:
            header_idx = i

    if header_idx is not None:
        lecturers: list[dict[str, str]] = []
        for row in rows[header_idx + 1 :]:
            entry = _lecturer_row_dict(row.cells)
            if any((entry.get(k) or "").strip() for k in entry):
                lecturers.append(entry)
        result["lecturers"] = lecturers
        if lecturers:
            result["lecturer"] = lecturers[0]
    return result


def _is_assessment_header(first: str, second: str) -> bool:
    return first in ("assessment", "assessment task") and "title and brief" in second


def _is_assessment_data_row(first: str) -> bool:
    if first.startswith("assessment task "):
        return len(first) > len("assessment task ")
    if first.startswith("assessment "):
        return len(first) > len("assessment ")
    return False


def _extract_assessments(doc: Document) -> list[dict[str, str]]:
    table = find_table(
        doc,
        lambda t: any(
            _is_assessment_header(
                norm(cell_text(distinct_cells(r)[0])),
                norm(cell_text(distinct_cells(r)[1])),
            )
            for r in t.rows
            if len(distinct_cells(r)) >= 2
        ),
    )
    assessments: list[dict[str, str]] = []
    if table is None:
        return assessments

    for row in table.rows:
        cells = distinct_cells(row)
        first = norm(cell_text(cells[0]))
        if not _is_assessment_data_row(first):
            continue
        assessments.append(
            {
                "number": cell_text(cells[0]),
                "title": cell_text(cells[1]) if len(cells) >= 2 else "",
                "due": cell_text(cells[2]) if len(cells) >= 3 else "",
            }
        )
    return assessments


def _extract_sessions(doc: Document) -> list[dict[str, str]]:
    table = find_table(
        doc,
        lambda t: any(
            norm(cell_text(r.cells[0])) == "session" and norm(cell_text(r.cells[1])) == "hrs"
            for r in t.rows
            if len(r.cells) >= 2
        ),
    )
    sessions: list[dict[str, str]] = []
    if table is None:
        return sessions

    rows = table.rows
    subheader = next(
        i
        for i, r in enumerate(rows)
        if len(r.cells) >= 2
        and norm(cell_text(r.cells[0])) == "session"
        and norm(cell_text(r.cells[1])) == "hrs"
    )
    for row in rows[subheader + 1 :]:
        c = row.cells
        first = norm(cell_text(c[0]))
        if first.startswith("total"):
            break
        get = lambda idx: cell_text(c[idx]) if idx < len(c) else ""
        sessions.append(
            {
                "session": get(0),
                "hrs": get(1),
                "element": get(2),
                "topic": get(3),
                "resources": get(4),
                "ooc_activity": get(5),
                "ooc_hrs": get(6),
            }
        )
    return sessions


def extract_lap(doc: Document | bytes | Path | str) -> dict:
    """Read LAP content from an old- or new-format document."""
    if isinstance(doc, (bytes, bytearray)):
        doc = Document(BytesIO(doc))
    elif not hasattr(doc, "tables"):
        doc = Document(doc)
    units, location = _extract_units(doc)
    supply = _extract_lecturer_and_supply(doc)
    return {
        "qualification": _info_value(doc, "qualification national code"),
        "delivery_period": _info_value(doc, "delivery period"),
        "cluster": _info_value(doc, "cluster name"),
        "units": units,
        "delivery_location": location,
        "student_supply": supply["student_supply"],
        "college_supply": supply["college_supply"],
        "lecturer": supply["lecturer"],
        "lecturers": supply["lecturers"],
        "assessments": _extract_assessments(doc),
        "sessions": _extract_sessions(doc),
    }


def _units_table(doc: Document):
    return find_table(
        doc,
        lambda t: any(
            norm(cell_text(distinct_cells(r)[0])) in ("national id", "national code")
            for r in t.rows
        ),
    )


def _fill_info(doc: Document, data: dict) -> None:
    for needles, value in (
        (("qualification national code",), data["qualification"]),
        (("delivery period",), data["delivery_period"]),
        (("cluster name",), data["cluster"]),
    ):
        if not value:
            continue
        for table in doc.tables:
            row = find_row(table, *needles)
            if row is not None:
                cells = distinct_cells(row)
                if len(cells) >= 2:
                    set_cell_text(cells[-1], value)
                break


def set_delivery_period(doc: Document, value: str) -> None:
    """Write the Delivery period field in the LAP header table."""
    text = (value or "").strip()
    if not text:
        return
    for table in doc.tables:
        row = find_row(table, "delivery period")
        if row is not None:
            cells = distinct_cells(row)
            if len(cells) >= 2:
                set_cell_text(cells[-1], text)
            break


def _fill_units(doc: Document, data: dict) -> None:
    table = _units_table(doc)
    if table is None or not data["units"]:
        return
    rows = list(table.rows)
    header_idx = next(
        i
        for i, r in enumerate(rows)
        if norm(cell_text(distinct_cells(r)[0])) in ("national id", "national code")
    )
    first_unit_row = rows[header_idx + 1]

    def fill_unit_row(row, unit: dict[str, str]) -> None:
        cells = distinct_cells(row)
        set_cell_text(cells[0], unit["code"])
        if len(cells) >= 2:
            set_cell_text(cells[1], unit["name"])

    fill_unit_row(first_unit_row, data["units"][0])
    prev = first_unit_row
    for unit in data["units"][1:]:
        prev = copy_row_after(table, prev)
        fill_unit_row(prev, unit)

    if data["delivery_location"]:
        loc_row = find_row(table, "delivery location")
        if loc_row is not None:
            cells = distinct_cells(loc_row)
            if len(cells) >= 2:
                set_cell_text(cells[-1], data["delivery_location"])


def _fill_supply_and_lecturer(doc: Document, data: dict) -> None:
    table = find_table(
        doc,
        lambda t: any(
            "student to supply" in norm(cell_text(distinct_cells(r)[0])) for r in t.rows
        ),
    )
    if table is None:
        return
    rows = list(table.rows)
    for row in rows:
        cells = distinct_cells(row)
        first = norm(cell_text(cells[0]))
        if "student to supply" in first and data["student_supply"]:
            set_cell_text(cells[0], data["student_supply"])
        elif "college to supply" in first and data["college_supply"]:
            set_cell_text(cells[0], data["college_supply"])

    lecturers = list(data.get("lecturers") or [])
    if not lecturers and any((data.get("lecturer") or {}).get(k) for k in ("name", "phone", "email")):
        lecturers = [data["lecturer"]]
    if not lecturers:
        return

    header_idx = next(
        i
        for i, row in enumerate(rows)
        if "lecturer name" in norm(cell_text(distinct_cells(row)[0]))
    )

    def fill_lecturer_row(row, lec: dict[str, str]) -> None:
        value_cells = row.cells
        for idx, key in enumerate(("name", "phone", "email", "contact_times", "campus")):
            value = (lec.get(key) or "").strip()
            if value and idx < len(value_cells):
                set_cell_text(value_cells[idx], value)

    first_data_row = rows[header_idx + 1]
    fill_lecturer_row(first_data_row, lecturers[0])
    prev = first_data_row
    for lec in lecturers[1:]:
        prev = copy_row_after(table, prev)
        fill_lecturer_row(prev, lec)


def _fill_assessments(doc: Document, data: dict) -> None:
    table = find_table(
        doc,
        lambda t: any(
            norm(cell_text(distinct_cells(r)[0])) == "assessment task"
            and "title and brief" in norm(cell_text(distinct_cells(r)[1]))
            for r in t.rows
            if len(distinct_cells(r)) >= 2
        ),
    )
    if table is None or not data["assessments"]:
        return
    task_rows = [
        r for r in table.rows if norm(cell_text(distinct_cells(r)[0])).startswith("assessment task ")
    ]
    if not task_rows:
        return

    def fill_assessment_row(row, item: dict[str, str]) -> None:
        cells = distinct_cells(row)
        if len(cells) >= 2 and item["title"]:
            set_cell_text(cells[1], item["title"])

    prev = task_rows[-1]
    for i, item in enumerate(data["assessments"]):
        if i < len(task_rows):
            fill_assessment_row(task_rows[i], item)
            prev = task_rows[i]
        else:
            prev = copy_row_after(table, prev)
            set_cell_text(distinct_cells(prev)[0], f"Assessment task {i + 1}")
            fill_assessment_row(prev, item)


def _fill_sessions(doc: Document, data: dict) -> None:
    table = find_table(
        doc,
        lambda t: any(
            norm(cell_text(r.cells[0])) == "session" and norm(cell_text(r.cells[1])) == "hrs"
            for r in t.rows
            if len(r.cells) >= 2
        ),
    )
    if table is None or not data["sessions"]:
        return
    rows = list(table.rows)
    subheader = next(
        i
        for i, r in enumerate(rows)
        if len(r.cells) >= 2
        and norm(cell_text(r.cells[0])) == "session"
        and norm(cell_text(r.cells[1])) == "hrs"
    )

    by_number: dict[str, object] = {}
    for r in rows[subheader + 1 :]:
        first_raw = cell_text(r.cells[0])
        first = norm(first_raw)
        if first.startswith("total"):
            break
        if first:
            by_number[_session_lookup_key(first_raw)] = r

    def fill_session_row(row, s: dict[str, str]) -> None:
        c = row.cells
        for idx, key in enumerate(
            ("session", "hrs", "element", "topic", "resources", "ooc_activity", "ooc_hrs")
        ):
            value = (s.get(key) or "").strip()
            if idx < len(c) and value:
                if key == "session" and value.isdigit():
                    value = str(int(value))
                set_cell_text(c[idx], value)

    prev = None
    for s in data["sessions"]:
        num = _session_lookup_key(s["session"])
        if num and num in by_number:
            row = by_number[num]
            fill_session_row(row, s)
            prev = row
        elif prev is not None:
            prev = copy_row_after(table, prev)
            fill_session_row(prev, s)
            if num:
                by_number[num] = prev

    update_session_hour_totals(doc)


def _parse_hours_value(text: str) -> float:
    cleaned = (text or "").strip().lower().replace("hrs", "").strip()
    if not cleaned:
        return 0.0
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _format_hours_sum(total: float) -> str:
    if total == int(total):
        return str(int(total))
    return str(total)


def update_session_hour_totals(doc: Document) -> None:
    """Sum in-class and out-of-class session hours into the table footer rows."""
    table = find_table(
        doc,
        lambda t: any(
            norm(cell_text(r.cells[0])) == "session" and norm(cell_text(r.cells[1])) == "hrs"
            for r in t.rows
            if len(r.cells) >= 2
        ),
    )
    if table is None:
        return

    rows = list(table.rows)
    subheader = next(
        (
            i
            for i, r in enumerate(rows)
            if len(r.cells) >= 2
            and norm(cell_text(r.cells[0])) == "session"
            and norm(cell_text(r.cells[1])) == "hrs"
        ),
        None,
    )
    if subheader is None:
        return

    in_class_total = 0.0
    ooc_total = 0.0
    for row in rows[subheader + 1 :]:
        first = norm(cell_text(row.cells[0]))
        if first.startswith("total"):
            break
        if len(row.cells) >= 2:
            in_class_total += _parse_hours_value(cell_text(row.cells[1]))
        if len(row.cells) >= 7:
            ooc_total += _parse_hours_value(cell_text(row.cells[6]))

    training_row = None
    unit_row = None
    for row in rows:
        first = norm(cell_text(row.cells[0]))
        if "total training hours" in first or first == "total hours":
            training_row = row
        elif "total amount of training for this unit" in first:
            unit_row = row

    if training_row is not None:
        training_cells = distinct_cells(training_row)
        if len(training_cells) >= 2:
            set_cell_text(training_cells[1], _format_hours_sum(in_class_total))
        if len(training_cells) >= 4:
            set_cell_text(training_cells[3], _format_hours_sum(ooc_total))

    if unit_row is not None:
        unit_cells = distinct_cells(unit_row)
        if len(unit_cells) >= 2:
            combined = in_class_total + ooc_total
            set_cell_text(unit_cells[1], f"{_format_hours_sum(combined)}hrs")


def fill_lap(data: dict, template: Path | bytes | Document) -> bytes:
    """Write extracted LAP data into the current blank template."""
    if hasattr(template, "tables"):
        doc = template
    elif isinstance(template, (bytes, bytearray)):
        doc = Document(BytesIO(template))
    else:
        doc = Document(template)
    _fill_info(doc, data)
    _fill_units(doc, data)
    _fill_supply_and_lecturer(doc, data)
    _fill_assessments(doc, data)
    _fill_sessions(doc, data)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def migrate_lap_to_template(
    source: bytes | Path | str,
    *,
    template_path: Path | None = None,
) -> bytes:
    """Extract content from an uploaded LAP and write it onto the current template."""
    data = extract_lap(source)
    template = template_path or LAP_TEMPLATE_PATH
    if not template.is_file():
        raise FileNotFoundError(f"LAP template not found: {template}")
    return fill_lap(data, template)
