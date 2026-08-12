"""Write a qualification family back out as a CSP-shaped .docx.

The point is the round trip: a qualification imported from a CSP document, then
split into stages by hand, comes back out as a document with one table per
stage — and re-importing that document reproduces the same classes. So this
writes precisely what ``csp_qualification_import`` reads and nothing it cannot
parse:

  * a title paragraph (the importer takes the first as the qualification name)
  * per stage, a "Stage N" heading paragraph, which the importer recognises as
    a stage boundary
  * a table whose header carries SIN and TPN, class name and hours in column 0
    ("Name | 2hrs"), unit code in column 2, with continuation rows leaving
    column 0 blank for a class spanning several units
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt

# Column 0 is read as the class cell and column 2 as the unit code; the header
# must contain a cell reading exactly "SIN" and one containing "TPN" or the
# importer will not recognise the table at all.
HEADER = ["Class", "SIN", "TPN"]


@dataclass
class ExportClass:
    name: str
    hours: float | None = None
    unit_codes: list[str] = field(default_factory=list)


@dataclass
class ExportStage:
    label: str
    classes: list[ExportClass] = field(default_factory=list)


def _class_cell(cls: ExportClass) -> str:
    """"Cyber Support | 2hrs" — the exact shape _parse_class_cell expects."""
    if cls.hours is None:
        return cls.name
    hours = int(cls.hours) if float(cls.hours).is_integer() else cls.hours
    return f"{cls.name} | {hours}hrs"


def write_csp_docx(path: str | Path, *, title: str, stages: list[ExportStage]) -> Path:
    path = Path(path)
    doc = Document()

    heading = doc.add_paragraph(title)
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    for index, stage in enumerate(stages, start=1):
        # "Stage 1" prefix keeps the heading parseable as a stage boundary;
        # the stage's own name follows for a human reader.
        label = stage.label.strip()
        doc.add_paragraph(f"Stage {index}" + (f" — {label}" if label else ""))

        table = doc.add_table(rows=1, cols=len(HEADER))
        table.style = "Table Grid"
        for col, text in enumerate(HEADER):
            table.rows[0].cells[col].text = text

        for cls in stage.classes:
            codes = cls.unit_codes or [""]
            for position, code in enumerate(codes):
                row = table.add_row()
                # Only the first row of a class names it. A blank class cell is
                # how the importer recognises a continuation row, which is what
                # makes multi-unit classes survive the round trip.
                row.cells[0].text = _class_cell(cls) if position == 0 else ""
                row.cells[1].text = ""
                row.cells[2].text = code

        doc.add_paragraph("")

    doc.save(path)
    return path
