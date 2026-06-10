"""Shared helpers for NMTAFE Learning & Assessment Plan (.docx) documents."""
from __future__ import annotations

from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

_BLACK = RGBColor(0, 0, 0)
_LAP_FONT_NAME = "Aptos"
_LAP_FONT_SIZE = Pt(12)


def iter_block_items(doc):
    """Yield the document body's paragraphs and tables in their real order."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def norm(text: str) -> str:
    """Normalise label text for comparison: lowercase, no colons, single spaces."""
    return " ".join((text or "").lower().replace(":", " ").split())


def cell_text(cell) -> str:
    return (cell.text or "").strip()


def distinct_cells(row):
    """Return a row's cells with horizontally-merged duplicates collapsed."""
    seen: set[int] = set()
    out = []
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            out.append(cell)
    return out


def find_table(doc, predicate):
    """Return the first table matching predicate(table), or None."""
    for table in doc.tables:
        if predicate(table):
            return table
    return None


def find_row(table, *needles):
    """Return the first row whose first cell contains all needles, else None."""
    needles_norm = [norm(n) for n in needles]
    for row in table.rows:
        first = norm(cell_text(distinct_cells(row)[0]))
        if all(n in first for n in needles_norm):
            return row
    return None


def _force_lap_font(run) -> None:
    """Written LAP values use Aptos 12 pt black, not styles from old uploads."""
    run.font.name = _LAP_FONT_NAME
    run.font.size = _LAP_FONT_SIZE
    run.font.color.rgb = _BLACK
    r_pr = run._element.rPr
    if r_pr is None:
        return
    color_el = run.font.color._element
    if color_el is not None:
        for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
            color_el.attrib.pop(qn(attr), None)
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            r_fonts.attrib.pop(qn(attr), None)


def set_cell_text(cell, text: str) -> None:
    """Replace a cell's contents with plain black text."""
    lines = (text or "").split("\n")
    first_para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    _set_paragraph_text(first_para, lines[0] if lines else "")
    for line in lines[1:]:
        new_para = cell.add_paragraph()
        _set_paragraph_text(new_para, line)


def _set_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        run = runs[0]
        run.text = text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        run = paragraph.add_run(text)
    _force_lap_font(run)


def copy_row_after(table, source_row):
    """Deep-copy a table row and insert it right after source_row."""
    import copy

    from docx.table import _Row

    new_tr = copy.deepcopy(source_row._tr)
    source_row._tr.addnext(new_tr)
    return _Row(new_tr, table)


def footer_codes(doc) -> str:
    """Collect footer text (used to detect LAP template code F122A14)."""
    chunks: list[str] = []
    for section in doc.sections:
        for para in section.footer.paragraphs:
            chunks.append(para.text)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
    return " ".join(chunks)
